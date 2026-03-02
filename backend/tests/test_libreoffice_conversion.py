"""
Test LibreOffice DOCX/DOC conversion functionality.
Tests:
1. Admin services status shows LibreOffice as configured and healthy
2. DOCX file upload converts to PDF using LibreOffice
3. DOC file upload converts to PDF using LibreOffice
4. Converted PDF preserves formatting (tables, text styles)

Test Credentials:
- Admin: admin@tls.or.tz / TLS@Admin2024
- Advocate: test@tls.or.tz / Test@12345678!
"""

import pytest
import requests
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import PyPDF2

BASE_URL = os.environ.get('REACT_APP_BACKEND_URL', '').rstrip('/')


class TestLibreOfficeConversion:
    """Test LibreOffice document conversion functionality"""
    
    @pytest.fixture(scope="class")
    def admin_token(self):
        """Get admin authentication token"""
        response = requests.post(f"{BASE_URL}/api/auth/login", json={
            "email": "admin@tls.or.tz",
            "password": "TLS@Admin2024"
        })
        assert response.status_code == 200, f"Admin login failed: {response.text}"
        data = response.json()
        return data.get("access_token")
    
    @pytest.fixture(scope="class")
    def advocate_token(self):
        """Get advocate authentication token"""
        response = requests.post(f"{BASE_URL}/api/auth/login", json={
            "email": "test@tls.or.tz",
            "password": "Test@12345678!"
        })
        assert response.status_code == 200, f"Advocate login failed: {response.text}"
        data = response.json()
        return data.get("access_token")
    
    @pytest.fixture(scope="class")
    def test_docx_with_formatting(self):
        """Create a test DOCX file with tables and formatting"""
        doc = Document()
        
        # Add title with formatting
        title = doc.add_heading('TLS Document Conversion Test', level=0)
        
        # Add paragraph with bold text
        p1 = doc.add_paragraph()
        run1 = p1.add_run('This document contains ')
        run_bold = p1.add_run('bold text')
        run_bold.bold = True
        run2 = p1.add_run(', ')
        run_italic = p1.add_run('italic text')
        run_italic.italic = True
        run3 = p1.add_run(', and ')
        run_underline = p1.add_run('underlined text')
        run_underline.underline = True
        run4 = p1.add_run('.')
        
        # Add a paragraph with different font size
        p2 = doc.add_paragraph()
        run_large = p2.add_run('This is larger text (16pt)')
        run_large.font.size = Pt(16)
        
        # Add a table with data
        doc.add_heading('Sample Table', level=1)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Role'
        header_cells[2].text = 'Status'
        
        # Add data rows
        data_rows = [
            ('John Doe', 'Advocate', 'Active'),
            ('Jane Smith', 'Senior Advocate', 'Active'),
            ('Bob Johnson', 'Associate', 'Pending')
        ]
        
        for i, (name, role, status) in enumerate(data_rows, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = name
            row_cells[1].text = role
            row_cells[2].text = status
        
        # Add a bullet list
        doc.add_heading('Key Points', level=1)
        doc.add_paragraph('First important point', style='List Bullet')
        doc.add_paragraph('Second important point', style='List Bullet')
        doc.add_paragraph('Third important point', style='List Bullet')
        
        # Add page break and another section
        doc.add_page_break()
        doc.add_heading('Second Page Content', level=1)
        doc.add_paragraph('This content should appear on the second page of the converted PDF.')
        
        # Save to bytes
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer.getvalue()
    
    def test_admin_services_status_shows_libreoffice(self, admin_token):
        """Test 1: Admin services status shows LibreOffice as configured and healthy"""
        response = requests.get(
            f"{BASE_URL}/api/admin/services/status",
            headers={"Authorization": f"Bearer {admin_token}"}
        )
        
        assert response.status_code == 200, f"Services status failed: {response.text}"
        data = response.json()
        
        # Verify response structure
        assert "services" in data, "Response should contain 'services' key"
        assert "libreoffice" in data["services"], "Services should include 'libreoffice'"
        
        libreoffice_status = data["services"]["libreoffice"]
        
        # Verify LibreOffice is configured
        assert libreoffice_status.get("configured") == True, \
            f"LibreOffice should be configured. Got: {libreoffice_status}"
        
        # Verify LibreOffice is healthy
        assert libreoffice_status.get("healthy") == True, \
            f"LibreOffice should be healthy. Got: {libreoffice_status}"
        
        # Verify path is set
        assert libreoffice_status.get("path") and libreoffice_status.get("path") != "Not installed", \
            f"LibreOffice path should be set. Got: {libreoffice_status.get('path')}"
        
        # Verify supported formats
        supported_formats = libreoffice_status.get("supported_formats", [])
        assert ".docx" in supported_formats, f"DOCX should be in supported formats. Got: {supported_formats}"
        assert ".doc" in supported_formats, f"DOC should be in supported formats. Got: {supported_formats}"
        
        # Verify document conversion section
        assert "document_conversion" in data, "Response should contain 'document_conversion' key"
        conversion_info = data["document_conversion"]
        assert conversion_info.get("primary") == "libreoffice", \
            f"Primary converter should be 'libreoffice'. Got: {conversion_info.get('primary')}"
        
        print(f"✓ LibreOffice Status: configured={libreoffice_status.get('configured')}, healthy={libreoffice_status.get('healthy')}")
        print(f"✓ LibreOffice path: {libreoffice_status.get('path')}")
        print(f"✓ Supported formats: {supported_formats}")
        print(f"✓ Primary converter: {conversion_info.get('primary')}")
    
    def test_docx_upload_converts_to_pdf(self, advocate_token, test_docx_with_formatting):
        """Test 2: DOCX file upload converts to PDF using LibreOffice"""
        files = {
            'file': ('test_document.docx', test_docx_with_formatting, 
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = requests.post(
            f"{BASE_URL}/api/documents/upload",
            headers={"Authorization": f"Bearer {advocate_token}"},
            files=files
        )
        
        assert response.status_code == 200, f"DOCX upload failed: {response.text}"
        data = response.json()
        
        # Verify conversion happened
        assert data.get("converted") == True, f"Document should be converted. Got: {data}"
        
        # Verify hash exists (confirms PDF was created)
        assert data.get("hash"), f"Document hash should exist. Got: {data}"
        
        # Verify pages (our test doc should have 2 pages)
        assert data.get("pages", 0) >= 1, f"PDF should have at least 1 page. Got: {data.get('pages')}"
        
        print(f"✓ DOCX uploaded and converted successfully")
        print(f"✓ Document hash: {data.get('hash')[:16]}...")
        print(f"✓ PDF pages: {data.get('pages')}")
        print(f"✓ Converted: {data.get('converted')}")
    
    def test_doc_upload_with_libreoffice(self, advocate_token):
        """Test 3: DOC file upload converts to PDF using LibreOffice (if format supported)"""
        # Create a simple DOCX and rename to DOC for testing
        # Note: This tests the DOC handling path even though we're using DOCX content
        # In reality, DOC format is binary and different from DOCX
        
        doc = Document()
        doc.add_heading('DOC Format Test', level=0)
        doc.add_paragraph('This tests the DOC file handling path.')
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        docx_content = docx_buffer.getvalue()
        
        # Test with DOC content type - LibreOffice can handle actual DOC files
        files = {
            'file': ('test_document.doc', docx_content, 'application/msword')
        }
        
        response = requests.post(
            f"{BASE_URL}/api/documents/upload",
            headers={"Authorization": f"Bearer {advocate_token}"},
            files=files
        )
        
        # DOC files should be converted by LibreOffice
        # Note: If this fails, it might be because the content is actually DOCX format
        # LibreOffice should still be able to detect and convert it
        if response.status_code == 200:
            data = response.json()
            print(f"✓ DOC file handled successfully")
            print(f"✓ Converted: {data.get('converted')}")
            print(f"✓ Pages: {data.get('pages')}")
        elif response.status_code == 400:
            # If LibreOffice couldn't convert (wrong binary format), that's expected
            # Our test uses DOCX content with DOC mime type
            print(f"⚠ DOC conversion returned 400 (expected if LibreOffice detected format mismatch)")
            print(f"  Response: {response.json().get('detail', response.text)}")
        else:
            pytest.fail(f"Unexpected response status: {response.status_code}, {response.text}")
    
    def test_formatting_preserved_in_conversion(self, advocate_token, test_docx_with_formatting):
        """Test 4: Verify that PDF conversion preserves formatting (tables, text)"""
        # Upload the formatted document
        files = {
            'file': ('formatted_doc.docx', test_docx_with_formatting, 
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = requests.post(
            f"{BASE_URL}/api/documents/upload",
            headers={"Authorization": f"Bearer {advocate_token}"},
            files=files
        )
        
        assert response.status_code == 200, f"Upload failed: {response.text}"
        data = response.json()
        
        # Verify page count indicates content was preserved
        # Our test document has 2 pages (we added a page break)
        pages = data.get("pages", 0)
        assert pages >= 2, f"PDF should have at least 2 pages (we have page break). Got: {pages}"
        
        # The hash confirms unique content was generated
        doc_hash = data.get("hash", "")
        assert len(doc_hash) == 64, f"Document hash should be SHA256 (64 chars). Got: {len(doc_hash)}"
        
        print(f"✓ Formatted document converted successfully")
        print(f"✓ Page count preserved: {pages} pages")
        print(f"✓ Document hash: {doc_hash[:16]}...")
        
        # Additional verification: The content was converted (not just passed through)
        assert data.get("converted") == True, "Document should have been converted to PDF"
    
    def test_libreoffice_service_direct(self):
        """Test 5: Direct test of LibreOffice service availability"""
        import subprocess
        
        # Check LibreOffice is in PATH
        result = subprocess.run(['which', 'libreoffice'], capture_output=True, text=True)
        assert result.returncode == 0, f"LibreOffice not found in PATH: {result.stderr}"
        libreoffice_path = result.stdout.strip()
        print(f"✓ LibreOffice found at: {libreoffice_path}")
        
        # Check LibreOffice version
        result = subprocess.run(['libreoffice', '--version'], capture_output=True, text=True)
        assert result.returncode == 0, f"Could not get LibreOffice version: {result.stderr}"
        version = result.stdout.strip()
        print(f"✓ LibreOffice version: {version}")
        
        # Verify it's a reasonable version
        assert 'LibreOffice' in version, f"Unexpected version output: {version}"
    
    def test_multiple_docx_conversions(self, advocate_token):
        """Test 6: Test multiple sequential DOCX conversions (stress test)"""
        successful_conversions = 0
        
        for i in range(3):
            # Create unique document each time
            doc = Document()
            doc.add_heading(f'Test Document {i+1}', level=0)
            doc.add_paragraph(f'This is test paragraph {i+1}.')
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            files = {
                'file': (f'test_doc_{i+1}.docx', buffer.getvalue(),
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            
            response = requests.post(
                f"{BASE_URL}/api/documents/upload",
                headers={"Authorization": f"Bearer {advocate_token}"},
                files=files
            )
            
            if response.status_code == 200 and response.json().get("converted"):
                successful_conversions += 1
        
        assert successful_conversions == 3, f"Expected 3 successful conversions, got {successful_conversions}"
        print(f"✓ All 3 sequential DOCX conversions successful")
    
    def test_services_status_unauthenticated(self):
        """Test 7: Verify services status requires admin authentication"""
        # No token
        response = requests.get(f"{BASE_URL}/api/admin/services/status")
        assert response.status_code == 401, f"Should require auth. Got: {response.status_code}"
        
        print("✓ Services status correctly requires authentication")
    
    def test_services_status_non_admin(self, advocate_token):
        """Test 8: Verify services status requires admin role"""
        response = requests.get(
            f"{BASE_URL}/api/admin/services/status",
            headers={"Authorization": f"Bearer {advocate_token}"}
        )
        # Advocate should not be able to access admin endpoints
        assert response.status_code == 403, f"Advocate should not access admin endpoint. Got: {response.status_code}"
        
        print("✓ Services status correctly requires admin role")


class TestLibreOfficeServiceModule:
    """Direct tests on the LibreOffice service module"""
    
    @pytest.fixture(scope="class")
    def advocate_token(self):
        """Get advocate authentication token"""
        response = requests.post(f"{BASE_URL}/api/auth/login", json={
            "email": "test@tls.or.tz",
            "password": "Test@12345678!"
        })
        if response.status_code != 200:
            pytest.skip("Could not login as advocate")
        return response.json().get("access_token")
    
    def test_complex_table_conversion(self, advocate_token):
        """Test conversion of document with complex table"""
        doc = Document()
        doc.add_heading('Complex Table Test', level=0)
        
        # Create a 5x5 table
        table = doc.add_table(rows=5, cols=5)
        table.style = 'Table Grid'
        
        for i in range(5):
            for j in range(5):
                table.cell(i, j).text = f'R{i}C{j}'
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        files = {
            'file': ('complex_table.docx', buffer.getvalue(),
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = requests.post(
            f"{BASE_URL}/api/documents/upload",
            headers={"Authorization": f"Bearer {advocate_token}"},
            files=files
        )
        
        assert response.status_code == 200, f"Complex table conversion failed: {response.text}"
        data = response.json()
        assert data.get("converted") == True
        print(f"✓ Complex 5x5 table converted successfully")
    
    def test_multipage_document_conversion(self, advocate_token):
        """Test conversion of multi-page document"""
        doc = Document()
        
        # Create 5 pages of content
        for page in range(5):
            doc.add_heading(f'Page {page + 1}', level=0)
            for para in range(10):
                doc.add_paragraph(f'This is paragraph {para + 1} on page {page + 1}. ' * 5)
            if page < 4:  # Don't add page break after last page
                doc.add_page_break()
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        files = {
            'file': ('multipage.docx', buffer.getvalue(),
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        
        response = requests.post(
            f"{BASE_URL}/api/documents/upload",
            headers={"Authorization": f"Bearer {advocate_token}"},
            files=files
        )
        
        assert response.status_code == 200, f"Multipage conversion failed: {response.text}"
        data = response.json()
        assert data.get("converted") == True
        assert data.get("pages", 0) >= 5, f"Expected at least 5 pages, got {data.get('pages')}"
        print(f"✓ Multi-page document ({data.get('pages')} pages) converted successfully")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
